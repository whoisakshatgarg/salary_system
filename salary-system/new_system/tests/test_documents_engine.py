"""Fidelity spec for the document engine and the eight templates.

Every kind that has a reference document is filled with **that document's own
data** and its ``format_spec.py`` dump is diffed against
``reference_specs/<file>.spec.json``: merged ranges, column widths, row
heights, page setup and the fonts/borders/alignment/fills of every static cell
must match exactly, and the values must come back identical - reproducing the
reference is what proves the pipeline.

The only permitted differences are enumerated in ``DEVIATIONS`` below and each
one is asserted explicitly, so a deviation can never creep in silently.
"""

import json
import sys
import tempfile
import unittest
from pathlib import Path

NEW_SYSTEM = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(NEW_SYSTEM))
REPO = NEW_SYSTEM.parents[1]
sys.path.insert(0, str(REPO))

import format_spec                                                   # noqa: E402
from backend.documents import dates, engine, registry                # noqa: E402

REFERENCE_SPECS = REPO / "reference_specs"
SAMPLES = NEW_SYSTEM.parent / "sample docs"
TEMPLATES = NEW_SYSTEM / "backend" / "documents" / "templates"

# --------------------------------------------------------------------------
# the complete list of permitted departures from the references
# --------------------------------------------------------------------------
DEVIATIONS = """
1. CONVENTIONS §8 typo fixes:
     'Total Quotation Vaue' -> 'Total Quotation Value'   (quotation)
     'nas been tested'      -> 'has been tested'         (test certificate)
     'Payment. Terms'       -> 'Payment Terms'           (acknowledgement)
2. Straight apostrophes in generated date strings ("30th Jun' 2025"), where
   the references carry Word's curly U+2019.
3. The acknowledgement gains an explicit print setup (portrait A4,
   fitToWidth=1, fitToHeight=0) - the BIFF .xls conversion has none at all.
4. The two Test Certificate pictures come back as openpyxl images; the
   reference spec lists none because openpyxl cannot read the reference's
   drawing container (that is the bug being fixed, not a regression).
5. The Buyer's-Order line is regenerated in the canonical
   'PO#### Dtd. DD.MM.YY, ...' form.  This is CONVENTIONS §8 typo-normalisation
   policy, confirmed by the owner: the reference's copy was hand-typed with
   irregular separators (double spaces, a missing comma, a stray '&') and the
   {po, date_iso} list stays the editable source of truth.
6. The BOM placeholder's two example rows (9 and 10) are normalised to the
   plain black, non-italic font of slot rows 11-32 - real BOM lines must not
   print grey (owner, 2026-08-21).  Font only; the grey PLACEHOLDER footnote
   at A40 is untouched.
"""


# --------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------

def spec_of(path):
    """format_spec dump, JSON-normalised so it compares against the .spec.json."""
    path = Path(path)
    raw = (format_spec.xlsx_spec(path) if path.suffix == ".xlsx"
           else format_spec.docx_spec(path))
    return json.loads(json.dumps(raw, default=str))


def reference(name):
    return json.loads((REFERENCE_SPECS / name).read_text())


def norm_runs(runs):
    """Merge adjacent runs that share formatting.

    Word splits runs wherever it feels like it and our templates author each
    token as one run, so run *boundaries* carry no meaning - run *formatting*
    does.  Comparing the merged form catches every formatting change while
    ignoring the split noise.
    """
    out = []
    for r in runs:
        key = (bool(r["bold"]), bool(r["italic"]), r["font"], r["size"])
        if out and out[-1][0] == key:
            out[-1][1] += r["text"]
        else:
            out.append([key, r["text"]])
    return [[k, t] for k, t in out if t]


def norm_paragraphs(paragraphs):
    return [{"style": p["style"], "align": p["align"], "runs": norm_runs(p["runs"])}
            for p in paragraphs]


def table_shape(table):
    return [[(c["colspan"], c["text"]) for c in row] for row in table["rows"]]


def _drawing_extents(xlsx_path):
    """``(cx, cy)`` EMU extents of every *picture* anchor, straight from the zip.

    Namespace-agnostic on purpose: Excel writes ``xdr:twoCellAnchor`` with the
    extent inside the picture's ``a:xfrm``, openpyxl writes an unprefixed
    ``oneCellAnchor`` with the extent on the anchor - the drawn size is the
    same number either way.
    """
    import xml.etree.ElementTree as ET
    import zipfile

    def tag(el):
        return el.tag.rsplit("}", 1)[-1]

    with zipfile.ZipFile(xlsx_path) as z:
        root = ET.fromstring(z.read("xl/drawings/drawing1.xml"))
    out = []
    for anchor in root:
        if not any(tag(e) == "pic" for e in anchor.iter()):
            continue                                   # freeform shapes etc.
        for el in anchor.iter():
            if tag(el) == "ext" and el.get("cx"):
                out.append((int(el.get("cx")), int(el.get("cy"))))
                break
    return out


class DocBase(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls._tmp = tempfile.TemporaryDirectory()
        cls.tmp = Path(cls._tmp.name)

    @classmethod
    def tearDownClass(cls):
        cls._tmp.cleanup()

    def render(self, kind, payload, name=None):
        data, filename = engine.render(kind, payload)
        path = self.tmp / (name or filename)
        path.write_bytes(data)
        return path, filename

    # -- xlsx comparison ---------------------------------------------------
    def assert_sheet_matches(self, got, ref, sheet, *, value_exceptions=(),
                             style_exceptions=(), page_setup_override=None,
                             ignore_images=False):
        g, r = got["sheets"][sheet], ref["sheets"][sheet]
        self.assertEqual(g["merged_ranges"], r["merged_ranges"], "merged_ranges")
        self.assertEqual(g["col_widths"], r["col_widths"], "col_widths")
        self.assertEqual(g["row_heights"], r["row_heights"], "row_heights")
        self.assertEqual(g["page_setup"], page_setup_override or r["page_setup"],
                         "page_setup")
        if not ignore_images:
            self.assertEqual(g["images"], r["images"], "images")

        gc, rc = g["cells"], r["cells"]
        self.assertEqual(sorted(set(rc) - set(gc)), [], "cells missing from output")
        self.assertEqual(sorted(set(gc) - set(rc)), [], "cells added by output")
        exceptions = set(value_exceptions)
        style_skip = set(style_exceptions)
        style_diffs, value_diffs = [], []
        for coord in sorted(rc):
            a, b = dict(gc[coord]), dict(rc[coord])
            av, bv = a.pop("value"), b.pop("value")
            if a != b and coord not in style_skip:
                style_diffs.append((coord, b, a))
            if av != bv and coord not in exceptions:
                value_diffs.append((coord, bv, av))
        self.assertEqual(style_diffs, [], "font/border/alignment/fill/numfmt drift")
        self.assertEqual(value_diffs, [], "unexpected value differences")

    def assert_values(self, got, sheet, expected):
        cells = got["sheets"][sheet]["cells"]
        actual = {c: cells.get(c, {}).get("value") for c in expected}
        self.assertEqual(actual, expected)


# ==========================================================================
# reference payloads - each is the reference document's own data
# ==========================================================================

WORK_ORDER_160_23 = {
    "number": "160/23", "order_no": "ORD-160", "wo_no_short": "160/23",
    "client_code": "S01", "cust_po": "PO00837", "date_iso": "2023-05-16",
    "customer": {"name": "SELCO Products Company", "country": "USA",
                 "address_lines": [], "contact": {}},
    "currency": {"code": "USD", "symbol": "$", "header": "Prices (U.S.D.)"},
    "items": [
        {"sno": 1, "part_no": "TPS-2501-S", "item": "Custom Housing", "qty": 100,
         "material": "304SS", "marking": "-", "remarks": None},
        {"sno": 2, "part_no": "TWS-L003571", "item": "Housing for Magnetic Pickup",
         "qty": 4, "material": "304SS", "marking": "-", "remarks": None},
        {"sno": 3, "part_no": "TWB02000750", "item": 'Thermowell, 1/4"NPT (P01)',
         "qty": 400, "material": "Brass", "marking": "-", "remarks": None},
    ],
}

TEST_CERT_047 = {
    "number": "AT/TC/59812/EI-047/24-25", "order_no": "ORD-47",
    "date_iso": "2024-05-10",
    "cert_no": "AT/TC/59812/EI-047/24-25", "cert_date_iso": "2024-05-10",
    "customer_line": "Reotemp Instrument Corporation, USA",
    "po": 59812, "po_date_iso": "2024-03-13",
    "invoice_no": "AT/EI/24-25/047", "invoice_date_iso": "2024-05-10",
    "extra_elements": [],
    "customer": {"name": "Reotemp Instrument Corporation", "country": "USA",
                 "address_lines": [], "contact": {}},
    "currency": {"code": "USD", "symbol": "$", "header": "Prices (U.S.D.)"},
    "items": [
        {"sno": 1, "item": "ST6316-1", "size": '1-3/8" Hex.', "qty": 100,
         "component": "-", "heat_no": "H4515", "material": "A479-316L",
         "chem": {"C": 0.022, "Mn": 1.32, "Si": 0.5, "P": 0.036, "S": 0.018,
                  "Cr": 16.89, "Ni": 10.51, "Mo": 2.18}},
        {"sno": 2, "item": "ST7.5316-H", "size": '9/8" Hex.', "qty": 150,
         "component": "-", "heat_no": "H4443", "material": "A479-316L",
         "chem": {"C": 0.019, "Mn": 1.48, "Si": 0.39, "P": 0.044, "S": 0.011,
                  "Cr": 16.66, "Ni": 10.05, "Mo": 2.02}},
        {"sno": 3, "item": "SW4316-S", "size": "28 ø", "qty": 100,
         "component": "-", "heat_no": "H3803", "material": "A479-316L",
         "chem": {"C": 0.028, "Mn": 1.52, "Si": 0.4, "P": 0.041, "S": 0.01,
                  "Cr": 16.6, "Ni": 10.06, "Mo": 2.05}},
    ],
}

BOM_PLACEHOLDER = {
    "number": "AT/BOM/26-27/___", "order_no": "ORD-1",
    "bom_no": "AT/BOM/26-27/___", "date_iso": None,
    "customer_line": None, "po": None, "po_date_iso": None,
    "wo_no": None, "part_assy": None,
    "customer": {"name": "", "country": "", "address_lines": [], "contact": {}},
    "currency": {"code": "USD", "symbol": "$", "header": "Prices (U.S.D.)"},
    "items": [
        {"sno": "1", "part_no": "TWB02000750",
         "description": 'Thermowell, 1/4"NPT (P01)', "size": '3/4" Hex.',
         "material": "Brass", "heat_or_os": "H4508", "source": "In-House",
         "qty_per": "1", "total_qty": "600", "unit": "EA",
         "remarks": "Example - overwrite"},
        {"sno": "2", "part_no": "PC-NPT-025", "description": 'Plastic Cap, 1/4"NPT',
         "size": '1/4" NPT', "material": "HDPE", "heat_or_os": "OS-0001",
         "source": "Outsourced - V01", "qty_per": "1", "total_qty": "600",
         "unit": "EA", "remarks": "Example - overwrite"},
    ],
}

ACK_E01 = {
    "number": "E01.04.08.26.01", "order_no": "ORD-2916", "date_iso": "2026-08-04",
    "bill_to_lines": ["East Coast Sensors", "20 Hathaway Drive",
                      "Stratford, CT 06615", "USA"],
    "ship_to_lines": ["East Coast Sensors", "20 Hathaway Drive",
                      "Stratford, CT 06615", "USA"],
    "cust_po": 2916, "po_date_iso": "2026-07-31",
    "quotation_ref": "E01/AT/030826/594", "client_code": "E01",
    "ack_ref": "E01.04.08.26.01", "ack_date_iso": "2026-08-04",
    "contacts": {"name": "Ed O'Neill", "email": "Ed@ecsensors.com",
                 "tel": "001-203-381-6267", "fax": "001-203-381-6278"},
    "price_basis": "Delivered by Fedex with another PO",
    "payment_terms": "N 30 by Wire Transfer",
    "ship_date_iso": "2026-09-16", "wo_no_long": "E01.04.08.26.252.26",
    "customer": {"name": "East Coast Sensors", "country": "USA",
                 "address_lines": ["20 Hathaway Drive", "Stratford, CT 06615"],
                 "contact": {"name": "Ed O'Neill", "email": "Ed@ecsensors.com",
                             "tel": "001-203-381-6267", "fax": "001-203-381-6278"}},
    "currency": {"code": "USD", "symbol": "$", "header": "Prices (U.S.D.)"},
    "items": [
        {"sno": 1, "code": "ECS-125-688-NB-BR",
         "description": 'Barstock Thermowell, OAL 5.25"',
         "material": "Naval Brass Grade 1", "qty": 30, "unit": "EA",
         "unit_price": 37.29, "total": 1118.7},
        {"sno": 2, "code": "ECS-223-242-NB-BR",
         "description": 'Barstock Thermowell, OAL 3.88"',
         "material": "Naval Brass Grade 1", "qty": 30, "unit": "EA",
         "unit_price": 29.57, "total": 887.1},
    ],
    "total": 2005.8,
}

COC_PO02940 = {
    "number": "COC-PO-02940-EI-122", "order_no": "ORD-2940",
    "date_iso": "2025-06-30",
    "customer_caps": "SELCO PRODUCTS COMPANY", "customer_short": "SELCO",
    "po": "PO02940", "invoice_no": "AT/EI/25-26/122",
    "invoice_date_iso": "2025-06-30", "part_desc": "TWS061000",
    "material": "ASTM A276 304L", "plating": "NA", "finishing": "NA",
    "qty_shipped": 100, "authenticator": "Q.A. MANAGER",
    "date_shipped_iso": "2025-06-30",
    "customer": {"name": "SELCO Products Company", "country": "USA",
                 "address_lines": [], "contact": {}},
    "currency": {"code": "USD", "symbol": "$", "header": "Prices (U.S.D.)"},
    "items": [],
}

_PL_HEADER = {
    "number": "AT/EI/26-27/168", "order_no": "ORD-168",
    "invoice_no": "AT/EI/26-27/168", "date_iso": "2026-08-14",
    "buyer_po_block": [
        {"po": "PO03864", "date_iso": "2026-03-06"},
        {"po": "PO03956", "date_iso": "2026-03-27"},
        {"po": "PO03969", "date_iso": "2026-04-06"},
        {"po": "PO04181", "date_iso": "2026-06-05"},
        {"po": "PO04205", "date_iso": "2026-06-15"},
        {"po": "PO04219", "date_iso": "2026-06-16"},
    ],
    "iec": "0509008631", "ad_code": "0292085 / 2690009",
    "consignee_lines": ["SELCO PRODUCTS COMPANY", "640, Maestro Drive, Suite 102\t",
                        "Reno, NV 89511, USA", "Tel No. :001-775-6745100 ",
                        "Contact: Cassie Halverson"],
    "buyer_lines": ["SELCO PRODUCTS COMPANY", "8780, Technology Way,",
                    "Reno, NV 89521-5908, USA\t", "Tel No. :001-775-6745100"],
    "pre_carriage": "", "place_receipt": "", "origin_country": "",
    "port_loading": "", "port_discharge": "NV, USA",
    "final_destination": "USA", "vessel": "AIR", "terms": "DDU",
    "marks_lines": ["1/3 TO 3/3", "AS ADDRESS", " 3 BOXES"],
    "hts_line": ("Parts & Accessories for Automatic Regulating & Controlling "
                 "Instruments – (H.T. Code) (9032.90.6080)"),
    "total_weight_line": "Total Weight in Kgs: Net Wt: 50.000       Gross Wt: 53.600",
    "customer": {"name": "SELCO Products Company", "country": "USA",
                 "address_lines": [], "contact": {}},
    "currency": {"code": "USD", "symbol": "$", "header": "Prices (U.S.D.)",
                 "head": "USD ($)"},
}

PACKING_LIST_168 = dict(_PL_HEADER, boxes=[
    {"box_label": "Box No. 1", "size": "16” x 11” x 10”",
     "net_wt": "20.600", "gross_wt": "21.800", "items": [
         {"sno": 1, "code_desc": "1.TPCR-236-0790,Heat Pump Sensor Probe For "
                                 "Temperature Sensor Devices, Matl. Copper ",
          "qty": 2000},
         {"sno": 2, "code_desc": "2.TPS-0353,Thermistor Probe Housing  For "
                                 "Temperature Sensor Devices, Matl. S/Steel ",
          "qty": 1300},
         {"sno": 3, "code_desc": "3.3200625, Custom Probe For Temperature Sensor "
                                 "Devices, Matl. Brass", "qty": 200},
     ]},
    {"box_label": "Box No. 2", "size": "16” x 11” x 10”",
     "net_wt": "17.700", "gross_wt": "18.900", "items": [
         {"sno": 1, "code_desc": "1. TWB3M4M0630, Custom Probe For Temperature "
                                 "Sensor Devices, Matl. Brass", "qty": 300},
     ]},
    {"box_label": "Box No. 3", "size": "12” x 10” x 9”",
     "net_wt": "11.700", "gross_wt": "12.900", "items": [
         {"sno": 1, "code_desc": "1. TW5535-S, Thermowell For Temperature Sensor "
                                 "Devices, Matl. 304S/Steel", "qty": 200},
         {"sno": 2, "code_desc": "2.EO234201,Thermowell ½”Npt For Temperature "
                                 "Sensor Devices, Matl. Brass ", "qty": 60},
     ]},
], totals={"qty": "4060", "net_wt": "50.000", "gross_wt": "53.600"})

INVOICE_168 = dict(_PL_HEADER,
                   origin_country="INDIA", place_receipt="GHAZIABAD",
                   port_loading="NEW DELHI",
                   gsp_line=registry.GSP_LINE,
                   amount_words=("USD Three  Thousand Eight  Hundred Fifty Eight  "
                                 "Only"),
                   totals={"qty": "4060", "net_wt": "50.000", "amount": "3858.00"},
                   items=[
                       {"sno": 1, "code_desc": "1.TPCR-236-0790,Heat Pump Sensor Probe "
                                               "For Temperature Sensor Devices, Matl. Copper ",
                        "po": "PO03864   ", "qty": 2000, "net_wt": "15.000",
                        "rate": "0.675", "amount": "1350.00"},
                       {"sno": 2, "code_desc": "2.EO234201,Thermowell ½”Npt For Temperature "
                                               "Sensor Devices, Matl. Brass (Replacements)",
                        "po": "PO03956   ", "qty": 60, "net_wt": "8.300",
                        "rate": "-", "amount": "-"},
                       {"sno": 3, "code_desc": "3.TPS-0353,Thermistor Probe Housing  For "
                                               "Temperature Sensor Devices, Matl. S/Steel ",
                        "po": "PO03969   ", "qty": 1300, "net_wt": "1.600",
                        "rate": "0.320", "amount": "416.00"},
                       {"sno": 4, "code_desc": "4.3200625, Custom Probe For Temperature "
                                               "Sensor Devices, Matl. Brass",
                        "po": "PO04181   ", "qty": 200, "net_wt": "3.600",
                        "rate": "3.820", "amount": "764.00"},
                       {"sno": 5, "code_desc": "5. TWB3M4M0630, Custom Probe For Temperature "
                                               "Sensor Devices, Matl. Brass",
                        "po": "PO04205   ", "qty": 300, "net_wt": "17.600",
                        "rate": "2.980", "amount": "894.00"},
                       {"sno": 6, "code_desc": "6. TW5535-S, Thermowell For Temperature "
                                               "Sensor Devices, Matl. 304S/Steel",
                        "po": "PO04219   ", "qty": 200, "net_wt": "3.900",
                        "rate": "2.170", "amount": "434.00"},
                   ])

QUOTATION_316 = {
    "number": "T04/AT/130826/316", "order_no": "ORD-316", "date_iso": "2026-08-13",
    "client_code": "T04", "rfq_ref": "Email", "rfq_date_iso": "2026-08-11",
    "customer": {"name": "Thermosense Ltd.", "country": "England",
                 "address_lines": [],
                 "contact": {"name": "Ryan Davis", "tel": "(+44) 1628 531166",
                             "fax": "(+44) 1628 531499",
                             "email": "ryan.davis@thermosense.co.uk"}},
    "currency": {"code": "GBP", "symbol": "£", "header": "Prices (G.B.P.)"},
    "items": [
        {"sno": 1, "code": "-",
         "description": ("Collar Assy., Collar - 12mm OD x 6.5mm ID x 6mm Long welded "
                         "to Tube - 4mm OD x 3.6mm (-0.0/+0.2) ID x 20mm Long, 304SS"),
         "qty": 120000, "unit": "EA", "unit_price": 0.534, "total": 64080.00},
    ],
    "total": 64080.00,
    "validity_line": "THIS QUOTATION IS VALID FOR 30 DAYS",
    "price_basis": "Delivered by Fedex or similar in lots of 20000 pcs.",
    "lead_time": "12 Weeks for 1st Lot, balance lots 4 months apart",
    "payment_terms": ("N 30 by Wire transfer to our bank account in India. All bank "
                      "charges in UK to Thermosense account."),
    "guarantee": "Items are Guaranteed for 12 months from despatch",
    "taxes_duties": "All taxes, customs duties and levies in country of delivery to client account",
    "note": "Our offer is subject to Delhi, India Jurisdiction.",
    "approved_by": "Sumesh Garg", "prepared_by": "",
}


# ==========================================================================
# xlsx kinds with a reference spec
# ==========================================================================

class WorkOrderFidelity(DocBase):
    """work_order.xlsx filled with WO 160/23 must reproduce the reference."""

    def setUp(self):
        self.path, self.filename = self.render("work_order", WORK_ORDER_160_23)
        self.got = spec_of(self.path)
        self.ref = reference("Apex_Work_Order_(1).spec.json")

    def test_structure_and_values_match_exactly(self):
        for sheet in ("Sheet1", "Sheet2", "Sheet3"):
            self.assert_sheet_matches(self.got, self.ref, sheet)

    def test_logo_survives(self):
        self.assertEqual(self.got["sheets"]["Sheet1"]["images"],
                         [{"anchor_col": 5, "anchor_row": 0, "width": 57,
                           "height": 58}])

    def test_date_is_a_real_date_with_the_template_number_format(self):
        cell = self.got["sheets"]["Sheet1"]["cells"]["G8"]
        self.assertEqual(cell["num_format"], "dd\\/mm\\/yyyy")
        self.assertEqual(cell["value"], "2023-05-16 00:00:00")

    def test_download_filename(self):
        self.assertEqual(self.filename, "Apex-Work-Order-160-23.xlsx")


class TestCertFidelity(DocBase):
    """test_cert.xlsx filled with certificate 047 must reproduce the reference."""

    def setUp(self):
        self.path, self.filename = self.render("test_cert", TEST_CERT_047)
        self.got = spec_of(self.path)
        self.ref = reference("Test_Certiticate_PO59812-EI-047.spec.json")

    def test_structure_and_values_match_except_the_typo_fix(self):
        self.assert_sheet_matches(self.got, self.ref, "Sheet2",
                                  value_exceptions=["A33"], ignore_images=True)

    def test_typo_is_fixed(self):
        ref_text = self.ref["sheets"]["Sheet2"]["cells"]["A33"]["value"]
        got_text = self.got["sheets"]["Sheet2"]["cells"]["A33"]["value"]
        self.assertIn("nas been tested", ref_text)
        self.assertEqual(got_text, ref_text.replace("nas been tested",
                                                    "has been tested"))

    def test_pictures_are_reanchored(self):
        """openpyxl drops the reference's drawing container; we put them back.

        ``format_spec`` reports an image's *natural* pixel size (that is what
        openpyxl's reader fills in), so the anchored extents are checked
        against the reference's own EMU values straight out of the zip.
        """
        self.assertEqual(self.ref["sheets"]["Sheet2"]["images"], [])
        self.assertEqual(self.got["sheets"]["Sheet2"]["images"], [
            {"anchor_col": 0, "anchor_row": 0, "width": 203, "height": 175},
            {"anchor_col": 10, "anchor_row": 0, "width": 225, "height": 115},
        ])
        self.assertEqual(_drawing_extents(self.path),
                         _drawing_extents(SAMPLES /
                                          "Test Certiticate PO59812-EI-047.xlsx"))
        self.assertEqual(_drawing_extents(self.path),
                         [(973455, 842010), (933450, 476250)])

    def test_both_test_certificates_share_the_same_pictures(self):
        import zipfile
        media = {}
        for name in ("Test Certiticate PO59812-EI-047.xlsx",
                     "Test Certiticate PO60543-EI-100.xlsx"):
            with zipfile.ZipFile(SAMPLES / name) as z:
                media[name] = (z.read("xl/media/image1.png"),
                               z.read("xl/media/image2.png"))
        self.assertEqual(*media.values())
        self.assertEqual(media["Test Certiticate PO59812-EI-047.xlsx"][0],
                         (TEMPLATES / "media" / "apex_logo.png").read_bytes())
        self.assertEqual(media["Test Certiticate PO59812-EI-047.xlsx"][1],
                         (TEMPLATES / "media" / "apex_iso.png").read_bytes())

    def test_spare_columns_default_to_dash_and_accept_extra_elements(self):
        cells = self.got["sheets"]["Sheet2"]["cells"]
        self.assertEqual([cells[f"{c}8"]["value"] for c in "PQRST"], ["-"] * 5)

        payload = dict(TEST_CERT_047, extra_elements=["Cu", "Ti"])
        payload["items"] = [dict(i) for i in TEST_CERT_047["items"]]
        payload["items"][0]["spare_1"] = 0.31
        path, _ = self.render("test_cert", payload, name="tc_extra.xlsx")
        cells = spec_of(path)["sheets"]["Sheet2"]["cells"]
        self.assertEqual([cells[f"{c}8"]["value"] for c in "PQRST"],
                         ["Cu", "Ti", "-", "-", "-"])
        self.assertEqual(cells["P9"]["value"], "0.31")

    def test_download_filename(self):
        self.assertEqual(self.filename, "Test Certificate PO59812-EI-047.xlsx")


class BomFidelity(DocBase):
    """bom.xlsx filled with the placeholder's two example rows."""

    # deviation 6: rows 9/10 lose the placeholder's grey italic example font
    EXAMPLE_CELLS = [f"{c}{r}" for r in (9, 10)
                     for c in ("A", "B", "C", "F", "G", "H", "J", "L", "N", "P", "Q")]

    def setUp(self):
        self.path, self.filename = self.render("bom", BOM_PLACEHOLDER)
        self.got = spec_of(self.path)
        self.ref = reference("Apex_BOM_Template_PLACEHOLDER.spec.json")

    def test_structure_and_values_match_exactly(self):
        self.assert_sheet_matches(self.got, self.ref, "BOM",
                                  style_exceptions=self.EXAMPLE_CELLS)

    def test_example_rows_are_normalised_to_the_slot_styling(self):
        """The only styling delta: rows 9/10 take row 11's plain black font."""
        got, ref = self.got["sheets"]["BOM"]["cells"], self.ref["sheets"]["BOM"]["cells"]
        for coord in self.EXAMPLE_CELLS:
            model = f"{coord[0]}11"
            # the reference really was grey italic here ...
            self.assertTrue(ref[coord]["font"]["italic"], coord)
            self.assertEqual(ref[coord]["font"]["color"], "FF808080", coord)
            # ... and we print it exactly like a plain slot row
            self.assertEqual(got[coord]["font"], ref[model]["font"], coord)
            self.assertFalse(got[coord]["font"]["italic"], coord)
            self.assertEqual(got[coord]["font"]["color"], "FF000000", coord)
            # font only - everything else still matches the reference
            for key in ("align", "borders", "fill", "num_format"):
                self.assertEqual(got[coord][key], ref[coord][key], f"{coord}.{key}")

    def test_no_other_cell_changed_styling(self):
        got, ref = self.got["sheets"]["BOM"]["cells"], self.ref["sheets"]["BOM"]["cells"]
        drifted = [c for c in ref
                   if {k: v for k, v in got[c].items() if k != "value"}
                   != {k: v for k, v in ref[c].items() if k != "value"}]
        self.assertEqual(sorted(drifted), sorted(self.EXAMPLE_CELLS))

    def test_placeholder_footnote_is_verbatim(self):
        self.assertEqual(self.got["sheets"]["BOM"]["cells"]["A40"]["value"],
                         self.ref["sheets"]["BOM"]["cells"]["A40"]["value"])

    def test_images_survive(self):
        self.assertEqual(self.got["sheets"]["BOM"]["images"],
                         self.ref["sheets"]["BOM"]["images"])
        self.assertEqual(len(self.got["sheets"]["BOM"]["images"]), 2)

    def test_download_filename(self):
        self.assertEqual(self.filename, "Apex-BOM-26-27.xlsx")   # placeholder no.
        self.assertEqual(registry.filename("bom", {"bom_no": "AT/BOM/26-27/001"}),
                         "Apex-BOM-26-27-001.xlsx")


class AckFidelity(DocBase):
    """ack.xlsx keeps the blank template's geometry and takes E01's data.

    The two real acknowledgements have different total-row positions because
    rows were inserted by hand over the years (CONVENTIONS §7), so the
    *structure* is diffed against the blank template's spec and the *values*
    against the filled E01 document, field by field.
    """

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls.tpl = reference("PO_Acknowledgement_Template_(converted).spec.json")
        cls.filled = reference("E01.04.08.26.252.26_(converted).spec.json")

    def setUp(self):
        tpl_cells = self.tpl["sheets"]["01"]["cells"]
        payload = dict(ACK_E01)
        # keep the blank template's own remittance wording + currency header so
        # the structural diff is clean; the 2026 defaults are tested separately
        payload.update(
            currency_header=tpl_cells["I18"]["value"],
            remit_intro=tpl_cells["A44"]["value"],
            remittance_block=tpl_cells["A45"]["value"],
            beneficiary_label=tpl_cells["A48"]["value"],
            beneficiary_address=tpl_cells["A49"]["value"],
            beneficiary_account=tpl_cells["A50"]["value"],
        )
        self.payload = payload
        self.path, self.filename = self.render("ack", payload)
        self.got = spec_of(self.path)

    def _variable_cells(self):
        entry = registry.entry("ack")
        cells = {s["cell"] for s in entry["tokens"].values()}
        items = entry["items"]
        for row in range(items["first_row"], items["last_row"] + 1):
            cells |= {f"{c}{row}" for c in items["cells"]}
        return cells

    def test_structure_matches_the_blank_template(self):
        self.assert_sheet_matches(
            self.got, self.tpl, "01",
            value_exceptions=self._variable_cells() | {"E15"},
            page_setup_override={"orientation": "portrait", "fitToWidth": 1,
                                 "fitToHeight": 0, "print_area": ""})

    def test_print_setup_is_the_documented_addition(self):
        self.assertEqual(self.tpl["sheets"]["01"]["page_setup"],
                         {"orientation": None, "fitToWidth": None,
                          "fitToHeight": None, "print_area": ""})
        self.assertEqual(self.got["sheets"]["01"]["page_setup"],
                         {"orientation": "portrait", "fitToWidth": 1,
                          "fitToHeight": 0, "print_area": ""})

    def test_payment_terms_typo_is_fixed(self):
        self.assertEqual(self.tpl["sheets"]["01"]["cells"]["E15"]["value"],
                         "Payment. Terms")
        self.assertEqual(self.got["sheets"]["01"]["cells"]["E15"]["value"],
                         "Payment Terms")

    def test_values_reproduce_the_filled_reference(self):
        ref = self.filled["sheets"]["01"]["cells"]
        self.assert_values(self.got, "01", {
            "A7": ref["A7"]["value"], "A8": ref["A8"]["value"],
            "A9": ref["A9"]["value"], "A10": ref["A10"]["value"],
            "G7": ref["G7"]["value"], "G8": ref["G8"]["value"],
            "G9": ref["G9"]["value"], "G10": ref["G10"]["value"],
            "F6": ref["F6"]["value"],                      # Cust. P.O.No. 2916
            "F7": ref["F7"]["value"],                      # 31st Jul' 2026
            "F9": ref["F9"]["value"],                      # quotation ref
            "C12": ref["C11"]["value"],                    # client code E01
            "F12": ref["F11"]["value"],                    # ackn. ref
            "F13": ref["F12"]["value"],                    # 04th Aug' 2026
            "F14": ref["F13"]["value"],                    # price basis
            "F15": ref["F14"]["value"],                    # payment terms
            "F16": ref["F15"]["value"],                    # 16th Sep' 2026
            "F17": ref["F16"]["value"],                    # work order no
            "B14": ref["B13"]["value"], "B15": ref["B14"]["value"],
            "B16": ref["B15"]["value"], "B17": ref["B16"]["value"],
            # items (reference rows 19/20 -> template slots 20/21)
            "A20": ref["A19"]["value"], "B20": ref["B19"]["value"],
            "C20": ref["C19"]["value"], "F20": ref["F19"]["value"],
            "G20": ref["G19"]["value"], "H20": ref["H19"]["value"],
            "I20": ref["I19"]["value"], "J20": ref["J19"]["value"],
            "A21": ref["A20"]["value"], "B21": ref["B20"]["value"],
            "C21": ref["C20"]["value"], "F21": ref["F20"]["value"],
            "G21": ref["G20"]["value"], "H21": ref["H20"]["value"],
            "I21": ref["I20"]["value"], "J21": ref["J20"]["value"],
            "I38": ref["I32"]["value"],                    # TOTAL P.O. VALUE
        })

    def test_unused_slots_are_blank_but_keep_their_styling(self):
        cells = self.got["sheets"]["01"]["cells"]
        tpl = self.tpl["sheets"]["01"]["cells"]
        for row in range(22, 38):
            for col in "ABCFGHIJ":
                coord = f"{col}{row}"
                if coord not in tpl:
                    continue
                self.assertIsNone(cells[coord]["value"], coord)
                self.assertEqual(cells[coord]["borders"], tpl[coord]["borders"], coord)
                self.assertEqual(cells[coord]["font"], tpl[coord]["font"], coord)

    def test_registry_defaults_are_the_current_2026_remittance(self):
        entry = registry.entry("ack")
        block = entry["tokens"]["remittance_block"]["default"]
        self.assertIn("SWIFT CODE : BOFAUS3N", block)
        self.assertNotIn("CHIPS ABA", block)                # the older variant
        self.assertNotIn("FED WIRE", block)
        self.assertIn("NEWYORK", block)                     # quoted verbatim, §8
        self.assertEqual(entry["tokens"]["currency_header"]["default"],
                         "Prices (U.S.D.)")

    def test_download_filename(self):
        self.assertEqual(self.filename, "E01.04.08.26.252.26.xlsx")


# ==========================================================================
# docx kinds with a reference spec
# ==========================================================================

class PackingListFidelity(DocBase):

    def setUp(self):
        self.path, self.filename = self.render("packing_list", PACKING_LIST_168)
        self.got = spec_of(self.path)
        self.ref = reference("Apex-Export_Packing_List-EI-168.spec.json")

    def test_sections_match(self):
        self.assertEqual(self.got["sections"], self.ref["sections"])

    def test_paragraphs_match(self):
        self.assertEqual(norm_paragraphs(self.got["paragraphs"]),
                         norm_paragraphs(self.ref["paragraphs"]))

    def test_table_topology_and_values_match(self):
        g, r = self.got["tables"][0], self.ref["tables"][0]
        self.assertEqual((g["n_rows"], g["n_cols"], g["style"]),
                         (r["n_rows"], r["n_cols"], r["style"]))
        got_rows, ref_rows = table_shape(g), table_shape(r)
        # r0c2 is the hand-typed buyer-PO line (documented deviation 5)
        got_rows[0][2] = ref_rows[0][2] = (3, "<buyer po line>")
        self.assertEqual(got_rows, ref_rows)

    def test_vertical_merge_groups_reproduce_the_boxes(self):
        """Box 1 spans 3 rows, box 2 one (no vMerge), box 3 two."""
        from docx import Document
        from docx.oxml.ns import qn
        rows = Document(str(self.path)).tables[0]._tbl.findall(qn("w:tr"))

        def vmerge(row_idx, cell_idx):
            tc = rows[row_idx].findall(qn("w:tc"))[cell_idx]
            tcPr = tc.find(qn("w:tcPr"))
            v = None if tcPr is None else tcPr.find(qn("w:vMerge"))
            return None if v is None else (v.get(qn("w:val")) or "continue")

        for cell in (3, 4, 5, 6):                       # box no / size / net / gross
            self.assertEqual(vmerge(8, cell), "restart")     # box 1, line 1
            self.assertEqual(vmerge(9, cell), "continue")
            self.assertEqual(vmerge(10, cell), "continue")
            self.assertIsNone(vmerge(11, cell))              # box 2: single line
            self.assertEqual(vmerge(12, cell), "restart")    # box 3, line 1
            self.assertEqual(vmerge(13, cell), "continue")
        for row in range(7, 14):                        # marks cell spans the lot
            self.assertEqual(vmerge(row, 0),
                             "restart" if row == 7 else "continue")

    def test_buyer_po_line_carries_every_po_and_date(self):
        line = self.got["tables"][0]["rows"][0][2]["text"]
        self.assertEqual(
            line.split("\n")[1],
            "PO03864   Dtd. 06.03.26, PO03956   Dtd. 27.03.26, "
            "PO03969   Dtd. 06.04.26, PO04181   Dtd. 05.06.26, "
            "PO04205   Dtd. 15.06.26, PO04219   Dtd. 16.06.26")

    def test_download_filename(self):
        self.assertEqual(self.filename, "Apex-Export Packing List-EI-168.docx")


class CocFidelity(DocBase):

    def setUp(self):
        self.path, self.filename = self.render("coc", COC_PO02940)
        self.got = spec_of(self.path)
        self.ref = reference("COC-PO-02940-EI-122.spec.json")

    def test_sections_match(self):
        self.assertEqual(self.got["sections"], self.ref["sections"])
        self.assertEqual(self.got["tables"], [])

    def test_paragraphs_match_modulo_the_apostrophe_normalisation(self):
        got = norm_paragraphs(self.got["paragraphs"])
        ref = norm_paragraphs(self.ref["paragraphs"])
        self.assertEqual(len(got), len(ref))
        for g, r in zip(got, ref):
            self.assertEqual(g["style"], r["style"])
            self.assertEqual(g["align"], r["align"])
            self.assertEqual([[k, t.replace("’", "'")] for k, t in g["runs"]],
                             [[k, t.replace("’", "'")] for k, t in r["runs"]])

    def test_generated_dates_use_a_straight_apostrophe(self):
        texts = ["".join(r["text"] for r in p["runs"]) for p in self.got["paragraphs"]]
        self.assertIn("Dtd. 30th Jun' 2025", texts[6])
        self.assertTrue(texts[13].endswith("30th Jun' 2025"))
        self.assertIn("AUTHENTICATOR’S NAME", texts[12])      # static label unchanged

    def test_tab_alignment_is_preserved(self):
        texts = ["".join(r["text"] for r in p["runs"]) for p in self.got["paragraphs"]]
        ref = ["".join(r["text"] for r in p["runs"]) for p in self.ref["paragraphs"]]
        for i in (5, 7, 8, 9, 10, 11, 12, 13):
            self.assertEqual(texts[i].count("\t"), ref[i].count("\t"), i)
            self.assertEqual(texts[i].split("\t")[0], ref[i].split("\t")[0], i)

    def test_download_filename(self):
        self.assertEqual(self.filename, "COC-PO-02940-EI-122.docx")


# ==========================================================================
# kinds without a reference spec - structural self-consistency
# ==========================================================================

class InvoiceStructure(DocBase):
    """No native reference (.doc is legacy) - pending owner visual sign-off."""

    def setUp(self):
        self.path, self.filename = self.render("invoice", INVOICE_168)
        self.got = spec_of(self.path)
        self.pl = reference("Apex-Export_Packing_List-EI-168.spec.json")

    def test_page_setup_matches_the_packing_list(self):
        self.assertEqual(self.got["sections"], self.pl["sections"])

    def test_company_header_matches_the_packing_list_except_the_title(self):
        got = norm_paragraphs(self.got["paragraphs"])
        ref = norm_paragraphs(self.pl["paragraphs"])
        self.assertEqual(got[:3], ref[:3])
        self.assertEqual("".join(t for _k, t in got[3]["runs"]), "INVOICE")
        self.assertEqual(got[3]["align"], ref[3]["align"])
        self.assertEqual([k for k, _t in got[3]["runs"]],
                         [k for k, _t in ref[3]["runs"]])

    def test_header_table_topology_matches_through_the_terms_row(self):
        g, r = table_shape(self.got["tables"][0]), table_shape(self.pl["tables"][0])
        self.assertEqual([[c for c, _t in row] for row in g[:6]],
                         [[c for c, _t in row] for row in r[:6]])

    def test_goods_grid(self):
        rows = table_shape(self.got["tables"][0])
        self.assertEqual([t for _c, t in rows[6]],
                         ["Marks & Nos. & No. & Kinds of Pkgs.", "Description of Goods",
                          "Order No.", "Qty. (Nos.)", "Net Weight (Kgs.)",
                          "Rate in USD ($)", "Amount in USD ($)"])
        self.assertEqual([c for c, _t in rows[6]], [1, 2, 1, 1, 2, 1, 1])
        self.assertEqual(self.got["tables"][0]["n_cols"], 9)
        # 7 header rows + HTS + 6 items + GSP + totals + words + signature
        self.assertEqual(self.got["tables"][0]["n_rows"], 18)

    def test_items_gsp_totals_and_words(self):
        rows = table_shape(self.got["tables"][0])
        self.assertEqual(rows[8][1][1], INVOICE_168["items"][0]["code_desc"])
        self.assertEqual([t for _c, t in rows[9][1:]],
                         [INVOICE_168["items"][1]["code_desc"],
                          "PO03956   ", "60", "8.300", "-", "-"])
        self.assertEqual(rows[13][1][1], INVOICE_168["items"][5]["code_desc"])
        self.assertEqual(rows[14][1][1], registry.GSP_LINE)
        self.assertEqual([t for _c, t in rows[15]],
                         [INVOICE_168["total_weight_line"], "Total", "4060",
                          "50.000", "", "3858.00"])
        self.assertEqual(rows[16][0][1],
                         "Amount Chargeable USD Three  Thousand Eight  Hundred "
                         "Fifty Eight  Only (in words)")
        self.assertEqual(rows[16][4][1], "Total in USD ($)")
        self.assertEqual(rows[16][5][1], "3858.00")
        self.assertEqual(rows[17][0][1], registry.DECLARATION)
        self.assertIn("AUTHORISED SIGNATORY", rows[17][1][1])

    def test_no_vertical_merges_outside_the_marks_column(self):
        from docx import Document
        from docx.oxml.ns import qn
        for row in Document(str(self.path)).tables[0]._tbl.findall(qn("w:tr")):
            for idx, tc in enumerate(row.findall(qn("w:tc"))):
                tcPr = tc.find(qn("w:tcPr"))
                v = None if tcPr is None else tcPr.find(qn("w:vMerge"))
                if v is not None:
                    self.assertIn(idx, (0, 2), "only the marks / terms cells merge")

    def test_download_filename(self):
        self.assertEqual(self.filename, "Apex-Export Invoice-EI-168.docx")


class QuotationStructure(DocBase):
    """Rebuilt from the PDF geometry - pending owner visual sign-off."""

    def setUp(self):
        self.path, self.filename = self.render("quotation", QUOTATION_316)
        self.got = spec_of(self.path)
        self.sheet = self.got["sheets"]["Quotation"]

    def test_page_setup(self):
        self.assertEqual(self.sheet["page_setup"],
                         {"orientation": "portrait", "fitToWidth": 1,
                          "fitToHeight": 0, "print_area": ""})

    def test_column_widths_hold_the_measured_pdf_proportions(self):
        from backend.documents.templates.build import build_quotation as bq
        widths = [self.sheet["col_widths"][c] for c in "ABCDEFG"]
        self.assertEqual(widths, [bq.pt_to_chars(w) for w in bq.COL_WIDTHS_PT])
        px = sum(w * bq.MDW_PX for w in widths) + 5 * len(widths)
        self.assertAlmostEqual(px * bq.PT_PER_PX, 481.5, delta=0.5)

    def test_typo_is_fixed(self):
        self.assertEqual(self.sheet["cells"]["A50"]["value"], "Total Quotation Value")

    def test_header_block(self):
        cells = self.sheet["cells"]
        self.assertEqual(cells["A1"]["value"], "QUOTATION")
        self.assertEqual(cells["A2"]["font"]["color"], "FF1F497D")
        self.assertIn("4167651", cells["A4"]["value"])          # quotation/ack phone
        self.assertEqual(cells["A6"]["value"], "Thermosense Ltd.")
        self.assertEqual(cells["A9"]["value"], "England")
        self.assertEqual(cells["F7"]["value"], "11th Aug' 2026")
        self.assertEqual(cells["F8"]["value"], "13th Aug' 2026")
        self.assertEqual(cells["F9"]["value"], "T04/AT/130826/316")
        self.assertEqual(cells["C10"]["value"], "T04")
        self.assertEqual(cells["F17"]["value"], "Prices (G.B.P.)")

    def test_item_row_number_formats(self):
        cells = self.sheet["cells"]
        self.assertEqual(cells["D19"]["value"], "120000")
        self.assertEqual(cells["D19"]["num_format"],
                         r"[>=10000000]##\,##\,##\,##0;[>=100000]##\,##\,##0;#,##0")
        self.assertEqual(cells["F19"]["num_format"], '"£"#,##0.000')
        self.assertEqual(cells["G19"]["num_format"], '"£"#,##0.00')
        self.assertEqual(cells["G50"]["num_format"], '"£"#,##0.00')

    def test_validity_line_sits_inside_the_table(self):
        cells = self.sheet["cells"]
        self.assertEqual(cells["C49"]["value"], "THIS QUOTATION IS VALID FOR 30 DAYS")
        self.assertEqual(cells["C49"]["align"][0], "center")
        self.assertEqual(cells["A49"]["borders"].get("right"), "thin")

    def test_footer_block(self):
        cells = self.sheet["cells"]
        self.assertEqual(cells["A59"]["value"], "Approved by : Sumesh Garg")
        self.assertEqual(cells["D60"]["value"], "Format : AT/QTN/EXP/01")
        self.assertEqual(cells["A56"]["value"], "Note")
        self.assertEqual(cells["C56"]["value"],
                         "Our offer is subject to Delhi, India Jurisdiction.")

    def test_outer_frame_is_medium_on_every_row(self):
        """openpyxl pushes a merged range's borders out from its top-left cell,
        so a merge reaching column G must not break the frame."""
        cells = self.sheet["cells"]
        for row in range(1, 61):
            self.assertEqual(cells[f"A{row}"]["borders"].get("left"), "medium", row)
            self.assertEqual(cells[f"G{row}"]["borders"].get("right"), "medium", row)
        for col in "ABCDEFG":
            self.assertEqual(cells[f"{col}1"]["borders"].get("top"), "medium", col)
            self.assertEqual(cells[f"{col}60"]["borders"].get("bottom"), "medium", col)

    def test_download_filename(self):
        self.assertEqual(self.filename, "Thermosense-Ltd-Quotation-316.xlsx")


# ==========================================================================
# engine behaviour
# ==========================================================================

class ItemOverflow(DocBase):
    """More items than slots: rows are cloned and the merges below survive."""

    def setUp(self):
        self.tpl = reference("PO_Acknowledgement_Template_(converted).spec.json")
        payload = dict(ACK_E01)
        payload["items"] = [
            {"sno": n, "code": f"CODE-{n:03d}", "description": f"Widget {n}",
             "material": "Brass", "qty": 10 * n, "unit": "EA",
             "unit_price": 1.5, "total": 15.0 * n}
            for n in range(1, 23)                        # 22 items into 18 slots
        ]
        self.extra = 22 - 18
        self.path, _ = self.render("ack", payload, name="ack_overflow.xlsx")
        self.got = spec_of(self.path)

    def test_rows_were_inserted(self):
        cells = self.got["sheets"]["01"]["cells"]
        self.assertEqual(cells["B37"]["value"], "CODE-018")
        self.assertEqual(cells["B41"]["value"], "CODE-022")
        self.assertEqual(cells["A41"]["value"], "22")

    def test_cloned_rows_carry_the_slot_styling_and_height(self):
        tpl = self.tpl["sheets"]["01"]
        cells, heights = self.got["sheets"]["01"]["cells"], self.got["sheets"]["01"]["row_heights"]
        for col in "ABCFGHIJ":
            if f"{col}37" not in tpl["cells"]:
                continue
            src, dst = tpl["cells"][f"{col}37"], cells[f"{col}41"]
            self.assertEqual(dst["borders"], src["borders"], col)
            self.assertEqual(dst["font"], src["font"], col)
            self.assertEqual(dst["num_format"], src["num_format"], col)
        self.assertEqual(heights["41"], tpl["row_heights"]["37"])

    def test_merges_below_the_insertion_point_moved(self):
        got = set(self.got["sheets"]["01"]["merged_ranges"])
        for before, after in (("A38:H38", "A42:H42"), ("I38:J38", "I42:J42"),
                              ("A39:J40", "A43:J44"), ("A45:J47", "A49:J51"),
                              ("G51:J51", "G55:J55")):
            self.assertIn(after, got, f"{before} should have moved to {after}")
            self.assertNotIn(before, got)
        # ranges above the insertion point are untouched
        for keep in ("A1:J1", "A11:D11", "F17:J17", "C37:E37"):
            self.assertIn(keep, got)

    def test_total_moved_with_the_total_row(self):
        cells = self.got["sheets"]["01"]["cells"]
        self.assertEqual(cells["A42"]["value"], "TOTAL P.O. VALUE")
        self.assertEqual(cells["I42"]["value"], "2005.8")
        self.assertEqual(cells["G55"]["value"], "FORMAT: AT/ACK/EXP/01")

    def test_merge_count_grows_only_by_the_cloned_rows_own_merge(self):
        """Each slot row carries one C:E merge, so N clones add exactly N."""
        self.assertEqual(len(self.got["sheets"]["01"]["merged_ranges"]),
                         len(self.tpl["sheets"]["01"]["merged_ranges"]) + self.extra)
        for row in range(38, 42):
            self.assertIn(f"C{row}:E{row}",
                          self.got["sheets"]["01"]["merged_ranges"])


class EngineContract(DocBase):

    PAYLOADS = {
        "quotation": QUOTATION_316, "ack": ACK_E01,
        "work_order": WORK_ORDER_160_23, "test_cert": TEST_CERT_047,
        "bom": BOM_PLACEHOLDER, "packing_list": PACKING_LIST_168,
        "invoice": INVOICE_168, "coc": COC_PO02940,
    }

    def test_every_kind_renders_and_leaves_no_token(self):
        self.assertEqual(sorted(self.PAYLOADS), registry.kinds())
        for kind, payload in sorted(self.PAYLOADS.items()):
            with self.subTest(kind=kind):
                data, filename = engine.render(kind, payload)
                self.assertEqual(data[:2], b"PK", "not an OOXML package")
                path = self.tmp / f"sweep_{kind}{Path(filename).suffix}"
                path.write_bytes(data)
                spec = spec_of(path)
                blob = json.dumps(spec)
                self.assertNotIn("{{", blob, f"{kind} left a token behind")

    def test_templates_and_registry_agree_on_every_coordinate(self):
        """XlsxFiller asserts each declared cell really carries its token."""
        for kind in registry.kinds():
            entry = registry.entry(kind)
            if entry["format"] != "xlsx":
                continue
            with self.subTest(kind=kind):
                engine.XlsxFiller(entry).fill(self.PAYLOADS[kind])

    def test_template_drift_is_caught(self):
        entry = dict(registry.entry("work_order"))
        entry["tokens"] = dict(entry["tokens"],
                               wo_no_short={"cell": "B7", "path": "wo_no_short"})
        with self.assertRaises(engine.TemplateError):
            engine.XlsxFiller(entry).fill(WORK_ORDER_160_23)

    def test_unknown_token_in_template_is_rejected(self):
        entry = dict(registry.entry("work_order"))
        entry["tokens"] = {k: v for k, v in entry["tokens"].items()
                           if k != "client_code"}
        with self.assertRaises(engine.TemplateError):
            engine.XlsxFiller(entry).fill(WORK_ORDER_160_23)

    def test_unknown_kind(self):
        with self.assertRaises(KeyError):
            registry.entry("delivery_note")

    def test_every_template_file_exists(self):
        for kind in registry.kinds():
            entry = registry.entry(kind)
            self.assertTrue((TEMPLATES / entry["template"]).exists(), kind)
        self.assertTrue((TEMPLATES / "README.md").exists())


class DateRenderers(unittest.TestCase):
    """CONVENTIONS §4 - each paper kind declares which renderer it uses."""

    def test_ordinal_apostrophe(self):
        for iso, want in (("2026-08-04", "04th Aug' 2026"),
                          ("2026-08-13", "13th Aug' 2026"),
                          ("2026-07-31", "31st Jul' 2026"),
                          ("2023-05-25", "25th May' 2023"),
                          ("2025-06-30", "30th Jun' 2025"),
                          ("2026-05-16", "16th May' 2026"),
                          ("2026-01-01", "01st Jan' 2026"),
                          ("2026-01-02", "02nd Jan' 2026"),
                          ("2026-01-03", "03rd Jan' 2026"),
                          ("2026-01-11", "11th Jan' 2026"),
                          ("2026-01-12", "12th Jan' 2026"),
                          ("2026-01-13", "13th Jan' 2026"),
                          ("2026-01-21", "21st Jan' 2026"),
                          ("2026-01-22", "22nd Jan' 2026"),
                          ("2026-01-23", "23rd Jan' 2026")):
            self.assertEqual(dates.ordinal_apostrophe(iso), want, iso)

    def test_other_renderers(self):
        self.assertEqual(dates.ddmmyyyy("2026-08-14"), "14/08/2026")
        self.assertEqual(dates.dtd_ddmmyy("2026-03-06"), "Dtd. 06.03.26")
        self.assertEqual(dates.us_mmddyy("2024-05-10"), "05-10-24")
        self.assertEqual(dates.date_value("2023-05-16").isoformat(), "2023-05-16")

    def test_empty_input(self):
        for fn in (dates.ordinal_apostrophe, dates.ddmmyyyy,
                   dates.dtd_ddmmyy, dates.us_mmddyy):
            self.assertEqual(fn(None), "")
            self.assertEqual(fn(""), "")
        self.assertIsNone(dates.date_value(None))

    def test_wave_ones_numbering_module_is_the_source_of_truth(self):
        """documents/dates.py must defer to backend.core.numbering when present."""
        from backend.core import numbering
        self.assertFalse(dates.LOCAL_FALLBACK)
        self.assertEqual(dates.ordinal_apostrophe("2026-08-04"),
                         numbering.ordinal_apostrophe("2026-08-04"))
        self.assertEqual(dates.us_mmddyy("2024-05-10"),
                         numbering.us_mmddyy("2024-05-10"))


if __name__ == "__main__":
    unittest.main()
