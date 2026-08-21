"""packing_list.docx  <-  "Apex-Export Packing List-EI-168.docx".

Provenance: the export packing list, a native .docx.  Its single 16x9 table
carries both the header block and the box-wise goods grid.  Changes:

1. every variable *value* becomes a single-run ``{{token}}`` - labels, tabs
   and surrounding whitespace are left exactly as the reference has them,
2. the six filled item rows collapse to ONE marker row (r8) that the engine
   deep-copies per item; the Box No./Box Size/Net/Gross cells carry the
   vertical-merge group for their box,
3. nothing else - company header paragraphs, exporter block, goods-grid
   headers and the signature row stay static.
"""

from __future__ import annotations

from docx import Document

from . import common, docxutil as dx

MARKER_ROW = 8              # the row the engine clones per item
DROP_ROWS = (9, 10, 11, 12, 13)   # the reference's other five filled item rows


def build():
    src = common.sample("Apex-Export Packing List-EI-168.docx")
    doc = Document(str(src))
    tbl = doc.tables[0]._tbl
    rows = dx.rows(tbl)

    # ---- r0: invoice no & date, buyer's order block -----------------------
    r0 = dx.cells(rows[0])
    dx.tokenize_para(dx.paragraphs(r0[1])[1], "invoice_no_date")
    dx.tokenize_para(dx.paragraphs(r0[2])[1], "buyer_po_block")

    # ---- r1: IEC / AD codes ----------------------------------------------
    r1 = dx.cells(rows[1])
    dx.tokenize_para(dx.paragraphs(r1[1])[0], "iec", old="0509008631")
    dx.tokenize_para(dx.paragraphs(r1[1])[1], "ad_code", old="0292085 / 2690009")

    # ---- r2: consignee / buyer blocks ------------------------------------
    r2 = dx.cells(rows[2])
    for i, p in enumerate(dx.paragraphs(r2[0])[1:], start=1):
        dx.tokenize_para(p, f"consignee_{i}")
    for i, p in enumerate(dx.paragraphs(r2[1])[1:], start=1):
        dx.tokenize_para(p, f"buyer_{i}")

    # ---- r3..r5: transport grid ------------------------------------------
    # the reference leaves four of these values blank, so those paragraphs have
    # no run at all: clone the formatting of a filled sibling ('AIR').
    r3, r4 = dx.cells(rows[3]), dx.cells(rows[4])
    donor = dx.paragraphs(r4[0])[1]
    dx.tokenize_para(dx.paragraphs(r3[0])[1], "pre_carriage", donor=donor)
    dx.tokenize_para(dx.paragraphs(r3[1])[1], "place_receipt", donor=donor)
    dx.tokenize_para(dx.paragraphs(r3[2])[1], "origin_country", donor=donor)
    dx.tokenize_para(dx.paragraphs(r3[3])[1], "country_final_destination", old="USA")

    dx.tokenize_para(dx.paragraphs(r4[0])[1], "vessel", old="AIR")
    dx.tokenize_para(dx.paragraphs(r4[1])[1], "port_loading", donor=donor)
    dx.tokenize_para(dx.paragraphs(r4[2])[1], "terms", old="DDU")

    r5 = dx.cells(rows[5])
    dx.tokenize_para(dx.paragraphs(r5[0])[1], "port_discharge", old="NV, USA")
    dx.tokenize_para(dx.paragraphs(r5[1])[1], "final_destination", old="USA")

    # ---- r7: marks block + HTS description line --------------------------
    r7 = dx.cells(rows[7])
    for i, p in enumerate(dx.paragraphs(r7[0])[1:], start=1):
        dx.tokenize_para(p, f"marks_{i}")
    dx.tokenize_para(dx.paragraphs(r7[1])[0], "hts_line")

    # ---- r8: the item marker row -----------------------------------------
    r8 = dx.cells(rows[MARKER_ROW])
    for idx, field in ((1, "code_desc"), (2, "qty"), (3, "box_label"),
                       (4, "box_size"), (5, "net_wt"), (6, "gross_wt")):
        dx.tokenize_para(dx.paragraphs(r8[idx])[0], f"item.{field}")

    # ---- r14: totals ------------------------------------------------------
    r14 = dx.cells(rows[14])
    dx.tokenize_para(dx.paragraphs(r14[0])[0], "total_weight_line")
    dx.tokenize_para(dx.paragraphs(r14[1])[0], "totals_qty")
    dx.tokenize_para(dx.paragraphs(r14[4])[0], "totals_net_wt")
    dx.tokenize_para(dx.paragraphs(r14[5])[0], "totals_gross_wt")

    # ---- drop the reference's remaining filled item rows ------------------
    for i in sorted(DROP_ROWS, reverse=True):
        dx.delete_row(rows[i])

    dst = common.out("packing_list.docx")
    doc.save(str(dst))
    return dst


if __name__ == "__main__":
    print(build())
