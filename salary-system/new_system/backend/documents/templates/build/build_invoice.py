"""invoice.docx  <-  built from the packing-list skeleton.

Provenance (confirmed with the owner 2026-08-21): the reference export invoice
is a legacy Word ``.doc`` that no pure-Python library can read, so there is no
``reference_specs`` entry for it.  The template is therefore derived from its
sibling document - the packing list - which shares the page setup, the company
header paragraphs and the header-table topology through the Terms row; the
goods grid below that is rebuilt from the ``textutil`` text dump of the .doc
(``phase0/dumps/invoice_textutil.txt``) plus the owner's render.

**Pending owner visual sign-off.**  Its test asserts structural
self-consistency, not equality with a reference spec.

Row plan of the single table (13 template rows):

    0-5   header block, identical to the packing list
    6     goods header: Marks | Description(2) | Order No. | Qty. (Nos.)
          | Net Weight (Kgs.)(2) | Rate in <cur> | Amount in <cur>
    7     HTS description line (marks cell starts its vertical merge here)
    8     ITEM MARKER row - cloned per item by the engine
    9     GSP duty line, centred in the description column
    10    Total Weight line (span 3) + Total / qty / net wt / amount
    11    Amount Chargeable <words> (in words) + Total in <cur> + amount
    12    declaration (left) + FOR APEX THERMOCON PVT. LTD. / signatory (right)
"""

from __future__ import annotations

from docx import Document

from . import build_packing_list, common, docxutil as dx

DECLARATION = ("WE DECLARE THAT THIS INVOICE SHOWS THE ACTUAL PRICE OF THE GOODS "
               "DESCRIBED AND THAT ALL PARTICULARS ARE TRUE AND CORRECT.")

GOODS_HEADERS = {
    2: "Order No.",
    3: "Qty. (Nos.)",
    4: "Net Weight (Kgs.)",
    5: "Rate in {{currency_head}}",
    6: "Amount in {{currency_head}}",
}

ITEM_TOKENS = {
    1: "{{item.code_desc}}",
    2: "{{item.po}}",
    3: "{{item.qty}}",
    4: "{{item.net_wt}}",
    5: "{{item.rate}}",
    6: "{{item.amount}}",
}


def _set(tc, text):
    ps = dx.paragraphs(tc)
    dx.set_para_text(ps[0], text)
    for extra in ps[1:]:
        tc.remove(extra)


def build():
    # the packing-list template is the skeleton: same section, same header rows
    pl = build_packing_list.build()
    doc = Document(str(pl))

    for p in dx.paragraphs(doc.element.body):
        if dx.para_text(p).strip() == "PACKING LIST":
            dx.set_para_text(p, "INVOICE")
            break
    else:                                                    # pragma: no cover
        raise AssertionError("packing-list title paragraph not found")

    tbl = doc.tables[0]._tbl
    rows = dx.rows(tbl)
    assert len(rows) == 11, len(rows)                        # 0-8 + totals + signature

    # ---- r6 goods header ---------------------------------------------------
    hdr = dx.cells(rows[6])
    for idx, text in GOODS_HEADERS.items():
        _set(hdr[idx], text)

    # ---- r8 item marker row ------------------------------------------------
    # the packing list merges its Box No./Size/Net/Gross cells vertically; the
    # invoice has no box columns, so only the marks cell keeps its merge.
    marker = rows[8]
    for idx, token in ITEM_TOKENS.items():
        _set(dx.cells(marker)[idx], token)
    for idx in range(1, len(dx.cells(marker))):
        dx.set_vmerge(dx.cells(marker)[idx], None)

    # ---- r9 GSP duty line (clone of the marker row) ------------------------
    gsp = dx.clone(marker)
    marker.addnext(gsp)
    gsp_cells = dx.cells(gsp)
    for idx in range(1, len(gsp_cells)):
        _set(gsp_cells[idx], "{{gsp_line}}" if idx == 1 else "")
        dx.set_vmerge(gsp_cells[idx], None)
    dx.set_vmerge(gsp_cells[0], "continue")

    # ---- totals row --------------------------------------------------------
    rows = dx.rows(tbl)
    totals = rows[10]
    tc = dx.cells(totals)
    _set(tc[0], "{{total_weight_line}}")
    _set(tc[1], "Total")
    _set(tc[2], "{{totals_qty}}")
    _set(tc[3], "{{totals_net_wt}}")
    _set(tc[4], "")
    _set(tc[5], "{{totals_amount}}")

    # ---- amount-in-words row (clone of the totals row) ---------------------
    words = dx.clone(totals)
    totals.addnext(words)
    wc = dx.cells(words)
    _set(wc[0], "Amount Chargeable {{amount_words}} (in words)")
    for idx in (1, 2, 3):
        _set(wc[idx], "")
    _set(wc[4], "Total in {{currency_head}}")
    _set(wc[5], "{{totals_amount_words_value}}")

    # ---- declaration + signature ------------------------------------------
    rows = dx.rows(tbl)
    sign = dx.cells(rows[-1])
    _set(sign[0], "{{declaration}}")

    dst = common.out("invoice.docx")
    doc.save(str(dst))
    return dst


if __name__ == "__main__":
    print(build())
