#!/usr/bin/env python3
"""Dump the complete structure + formatting of an .xlsx or .docx file to JSON.

Usage:  python format_spec.py <file> [out.json]

Pure Python (openpyxl / python-docx). No LibreOffice, no system packages.
Run it on a reference file and on a generated file, then diff the two JSONs:
matching output = matching format (merges, widths, heights, fonts, borders,
fills, alignment, page setup).
"""
import json, sys
from pathlib import Path


def _color(c):
    try:
        if c is None:
            return None
        if getattr(c, "type", None) == "theme":
            return f"theme:{c.theme}"
        v = c.rgb
        return v if isinstance(v, str) else None
    except Exception:
        return None


def xlsx_spec(path):
    from openpyxl import load_workbook
    wb = load_workbook(path)
    out = {"file": Path(path).name, "type": "xlsx", "sheets": {}}
    for ws in wb.worksheets:
        cells = {}
        for row in ws.iter_rows():
            for c in row:
                borders = {s: getattr(c.border, s).style
                           for s in ("left", "right", "top", "bottom")
                           if getattr(c.border, s).style}
                fill = _color(c.fill.fgColor) if c.fill.patternType else None
                if c.value is None and not borders and not fill:
                    continue
                cells[c.coordinate] = {
                    "value": None if c.value is None else str(c.value),
                    "font": {"name": c.font.name, "size": c.font.size,
                             "bold": bool(c.font.bold), "italic": bool(c.font.italic),
                             "color": _color(c.font.color)},
                    "align": [c.alignment.horizontal, c.alignment.vertical,
                              bool(c.alignment.wrap_text)],
                    "borders": borders,
                    "fill": fill,
                    "num_format": None if c.number_format == "General" else c.number_format,
                }
        images = []
        for img in getattr(ws, "_images", []):
            try:
                frm = img.anchor._from
                images.append({"anchor_col": frm.col, "anchor_row": frm.row,
                               "width": img.width, "height": img.height})
            except Exception:
                images.append({"anchor": "unknown"})
        out["sheets"][ws.title] = {
            "merged_ranges": sorted(str(r) for r in ws.merged_cells.ranges),
            "col_widths": {k: round(v.width, 2) for k, v in ws.column_dimensions.items() if v.width},
            "row_heights": {int(k): v.height for k, v in ws.row_dimensions.items() if v.height},
            "images": images,
            "page_setup": {"orientation": ws.page_setup.orientation,
                           "fitToWidth": ws.page_setup.fitToWidth,
                           "fitToHeight": ws.page_setup.fitToHeight,
                           "print_area": ws.print_area},
            "cells": cells,
        }
    return out


def docx_spec(path):
    from docx import Document
    d = Document(path)

    def runs(p):
        return [{"text": r.text, "bold": r.bold, "italic": r.italic,
                 "font": r.font.name,
                 "size": r.font.size.pt if r.font.size else None}
                for r in p.runs if r.text]

    paragraphs = [{"style": p.style.name,
                   "align": str(p.alignment),
                   "runs": runs(p)}
                  for p in d.paragraphs if p.text.strip()]

    tables = []
    for t in d.tables:
        rows = []
        for r in t.rows:
            row, seen = [], None
            for c in r.cells:
                if c._tc is seen:            # same underlying element = merged span
                    row[-1]["colspan"] += 1
                else:
                    row.append({"text": c.text, "colspan": 1})
                    seen = c._tc
            rows.append(row)
        tables.append({"n_rows": len(t.rows), "n_cols": len(t.columns),
                       "style": t.style.name if t.style else None, "rows": rows})

    sections = [{"orientation": str(s.orientation),
                 "page_width_pt": s.page_width.pt if s.page_width else None,
                 "page_height_pt": s.page_height.pt if s.page_height else None,
                 "margins_pt": [s.left_margin.pt, s.right_margin.pt,
                                s.top_margin.pt, s.bottom_margin.pt]}
                for s in d.sections]
    return {"file": Path(path).name, "type": "docx",
            "paragraphs": paragraphs, "tables": tables, "sections": sections}


def main():
    if len(sys.argv) < 2:
        sys.exit(__doc__)
    src = Path(sys.argv[1])
    ext = src.suffix.lower()
    if ext == ".xlsx":
        spec = xlsx_spec(src)
    elif ext == ".docx":
        spec = docx_spec(src)
    else:
        sys.exit(f"Unsupported extension {ext!r}: convert legacy .doc/.xls to "
                 ".docx/.xlsx once (already done for this project), then rerun.")
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix(src.suffix + ".spec.json")
    dst.write_text(json.dumps(spec, indent=1, default=str))
    n = sum(len(s.get("cells", {})) for s in spec.get("sheets", {}).values()) if spec["type"] == "xlsx" else len(spec["tables"])
    print(f"wrote {dst}  ({'cells' if spec['type']=='xlsx' else 'tables'}: {n})")


if __name__ == "__main__":
    main()
